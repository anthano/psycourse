from __future__ import annotations

import itertools
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pytask import task

from psycourse.config import BLD_RESULTS

# ── Directory shortcuts ────────────────────────────────────────────────────────
_UNI = BLD_RESULTS / "univariate" / "continuous_analysis"
_LIP = _UNI / "lipid"
_PRS = _UNI / "prs"
_LPAN = _LIP / "panss"
_PPAN = _PRS / "panss"
_MEDADJ = _LIP / "medication_adjusted"
_DESC = BLD_RESULTS / "descriptive_stats"
_MED = BLD_RESULTS / "mediation_analysis"
_R2 = BLD_RESULTS / "incremental_r2"
_TAB = BLD_RESULTS / "tables"

# ── Table-of-contents definition ──────────────────────────────────────────────
# (sheet_name, full description, analysis group)
_TOC_ROWS = [
    # ── Reference ──────────────────────────────────────────────────────────────
    (
        "02_lipid_list",
        "Full list of analysed lipid species with class annotations",
        "Reference",
    ),
    (
        "03_n_lipid",
        "Sample size (N) per lipid regression analysis and covariate model",
        "Reference",
    ),
    (
        "04_n_prs",
        "Sample size (N) per PRS regression analysis and covariate model",
        "Reference",
    ),
    # ── PRS ~ Subtype Probability ───────────────────────────────────────────────
    (
        "05_prs_std",
        "PRS associations with severe psychosis subtype probability "
        "(standard covariates: age, sex, ancestry PCs)",
        "PRS ~ Subtype Probability",
    ),
    # ── Lipid Species ~ Subtype Probability (PANSS as covariate, sensitivity) ──
    (
        "06_lip_std",
        "Lipid Species associations with severe psychosis subtype probability "
        "(standard covariates: sex, BMI, illness duration, smoking status)",
        "Lipid Species ~ Subtype Probability",
    ),
    (
        "07_lip_cov_diag",
        "Lipid Species ~ subtype probability: sensitivity additionally controlling "
        "for psychiatric diagnosis",
        "Lipid Species ~ Subtype Probability",
    ),
    (
        "08_lip_cov_med",
        "Lipid Species ~ subtype probability: sensitivity additionally controlling "
        "for antipsychotic medication",
        "Lipid Species ~ Subtype Probability",
    ),
    (
        "09_lip_cov_med_diag",
        "Lipid Species ~ subtype probability: sensitivity additionally controlling "
        "for medication and diagnosis",
        "Lipid Species ~ Subtype Probability",
    ),
    (
        "10_lip_cov_panss_pos",
        "Lipid Species ~ subtype probability: sensitivity additionally controlling "
        "for PANSS Positive subscale score",
        "Lipid Species ~ Subtype Probability",
    ),
    (
        "11_lip_cov_panss_neg",
        "Lipid Species ~ subtype probability: sensitivity additionally controlling "
        "for PANSS Negative subscale score",
        "Lipid Species ~ Subtype Probability",
    ),
    (
        "12_lip_cov_panss_gen",
        "Lipid Species ~ subtype probability: sensitivity additionally controlling "
        "for PANSS General Psychopathology subscale score",
        "Lipid Species ~ Subtype Probability",
    ),
    (
        "13_lip_cov_panss_tot",
        "Lipid Species ~ subtype probability: sensitivity additionally controlling "
        "for PANSS Total Score",
        "Lipid Species ~ Subtype Probability",
    ),
    # ── Lipid Class Enrichment (PANSS as covariate, sensitivity) ───────────────
    (
        "14_enrich_std",
        "Lipid class enrichment analysis for severe psychosis subtype "
        "probability (standard covariates; permutation-based enrichment scores)",
        "Lipid Class Enrichment ~ Subtype Probability",
    ),
    (
        "15_enrich_cov_diag",
        "Lipid class enrichment: sensitivity additionally controlling "
        "for psychiatric diagnosis",
        "Lipid Class Enrichment ~ Subtype Probability",
    ),
    (
        "16_enrich_cov_med",
        "Lipid class enrichment: sensitivity additionally controlling "
        "for antipsychotic medication",
        "Lipid Class Enrichment ~ Subtype Probability",
    ),
    (
        "17_enrich_cov_med_diag",
        "Lipid class enrichment: sensitivity additionally controlling "
        "for medication and diagnosis",
        "Lipid Class Enrichment ~ Subtype Probability",
    ),
    (
        "18_enrich_cov_panss_pos",
        "Lipid class enrichment: sensitivity additionally controlling "
        "for PANSS Positive subscale score",
        "Lipid Class Enrichment ~ Subtype Probability",
    ),
    (
        "19_enrich_cov_panss_neg",
        "Lipid class enrichment: sensitivity additionally controlling "
        "for PANSS Negative subscale score",
        "Lipid Class Enrichment ~ Subtype Probability",
    ),
    (
        "20_enrich_cov_panss_gen",
        "Lipid class enrichment: sensitivity additionally controlling "
        "for PANSS General Psychopathology subscale score",
        "Lipid Class Enrichment ~ Subtype Probability",
    ),
    (
        "21_enrich_cov_panss_tot",
        "Lipid class enrichment: sensitivity additionally controlling "
        "for PANSS Total Score",
        "Lipid Class Enrichment ~ Subtype Probability",
    ),
    (
        "22_enrich_cov_antidepressants",
        "Lipid class enrichment: sensitivity additionally controlling "
        "for antidepressant medication",
        "Lipid Class Enrichment ~ Subtype Probability",
    ),
    (
        "23_enrich_cov_antipsychotics",
        "Lipid class enrichment: sensitivity additionally controlling "
        "for antipsychotic medication class",
        "Lipid Class Enrichment ~ Subtype Probability",
    ),
    (
        "24_enrich_cov_mood_stabilizers",
        "Lipid class enrichment: sensitivity additionally controlling "
        "for mood stabilizer medication",
        "Lipid Class Enrichment ~ Subtype Probability",
    ),
    (
        "25_enrich_cov_tranquilizers",
        "Lipid class enrichment: sensitivity additionally controlling "
        "for tranquilizer medication",
        "Lipid Class Enrichment ~ Subtype Probability",
    ),
    # ── PRS ~ PANSS subscales as outcome ────────────────────────────────────────
    (
        "26_prs_panss_std",
        "PRS associations with PANSS subscale scores as continuous outcomes "
        "(standard covariates; all four subscales stacked — see "
        "'panss_outcome' column: "
        "PANSS Positive, Negative, General, Total Score)",
        "PRS ~ PANSS Subscales as Outcome",
    ),
    # ── Lipid (Top 20) ~ PANSS subscales as outcome ────────────────────────────
    (
        "27_lip_panss_std",
        "Top-20 lipid associations with PANSS subscale scores as continuous "
        "outcomes (standard covariates; all four subscales stacked — see "
        "'panss_outcome' column)",
        "Lipid ~ PANSS Subscales as Outcome",
    ),
    (
        "28_lip_panss_cov_med",
        "Top-20 lipids ~ PANSS subscales as outcome: sensitivity additionally "
        "controlling for antipsychotic medication (all four subscales stacked)",
        "Lipid ~ PANSS Subscales as Outcome",
    ),
    (
        "29_lip_panss_cov_diag",
        "Top-20 lipids ~ PANSS subscales as outcome: sensitivity additionally "
        "controlling for psychiatric diagnosis (all four subscales stacked)",
        "Lipid ~ PANSS Subscales as Outcome",
    ),
    (
        "30_lip_panss_cov_med_diag",
        "Top-20 lipids ~ PANSS subscales as outcome: sensitivity additionally "
        "controlling for medication and diagnosis (all four subscales stacked)",
        "Lipid ~ PANSS Subscales as Outcome",
    ),
    # ── Lipid Enrichment ~ PANSS subscales as outcome ──────────────────────────
    (
        "31_enrich_panss_std",
        "Lipid class enrichment with PANSS subscale scores as outcomes "
        "(standard covariates; all available subscales stacked — see "
        "'panss_outcome' column: Positive, Negative, General, Total Score)",
        "Lipid Class Enrichment ~ PANSS as Outcome",
    ),
    # ── Incremental R² ─────────────────────────────────────────────────────────
    (
        "32_r2_incremental",
        "Incremental R² decomposition: variance in severe psychosis subtype "
        "probability explained by PRS block and lipid class block "
        "(permutation-based p-values; 20 000 permutations)",
        "Incremental R²",
    ),
    (
        "33_r2_individual",
        "Individual predictor ΔR²: unique variance explained by each "
        "significant PRS and each enriched lipid class score "
        "(permutation-based p-values; 20 000 permutations)",
        "Incremental R²",
    ),
    # ── Mediation ──────────────────────────────────────────────────────────────
    (
        "34_mediation",
        "Mediation analysis results: indirect effects of PRS on severe "
        "psychosis subtype probability via lipid species",
        "Mediation Analysis",
    ),
    # ── CCA ────────────────────────────────────────────────────────────────────
    (
        "35_cca_loadings",
        "Full CCA canonical loadings for the first canonical variate: "
        "PRS block (13 polygenic scores) and lipid class block (16 lipid classes). "
        "The 'block' column identifies the input block; 'variable' is the "
        "feature name; "
        "'loading' is the correlation between the feature and its "
        "canonical component score.",
        "CCA Analysis",
    ),
]


# ── docx TOC writer ───────────────────────────────────────────────────────────


def _set_cell_bg(cell, hex_color: str) -> None:
    """Set the background fill of a table cell (python-docx has no direct API)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _write_toc_docx(path: Path) -> None:
    """Write a formatted Word document containing the supplementary table TOC."""
    doc = Document()

    # ── Page margins (narrower for more content width) ────────────────────────
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # ── Title block ───────────────────────────────────────────────────────────
    title = doc.add_heading("Supplementary Tables", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading("Table of Contents", level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    note = doc.add_paragraph(
        f"Generated: {date.today().isoformat()}   ·   "
        "All analyses target the severe psychosis subtype probability unless "
        "stated otherwise.   ·   "
        "Sheets containing results from multiple PANSS subscales include a "
        "'panss_outcome' column that identifies the subscale."
    )
    note.style.font.size = Pt(9)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # spacer

    # ── Build table ───────────────────────────────────────────────────────────
    # Columns: #  |  Sheet  |  Description
    # Group-header rows span all columns and are shaded.
    HDR_COLOR = "2E4057"  # dark slate — header row
    GRP_COLOR = "D6E4F0"  # light blue — group divider rows
    ALT_COLOR = "F7FBFF"  # very light blue — alternate data rows
    WHITE = "FFFFFF"

    col_widths = [Inches(0.45), Inches(1.55), Inches(4.5)]

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"

    # Header row
    hdr_cells = tbl.rows[0].cells
    for cell, txt, w in zip(
        hdr_cells,
        ["#", "Sheet", "Description"],
        col_widths,
        strict=False,
    ):
        cell.text = txt
        cell.width = w
        _set_cell_bg(cell, HDR_COLOR)
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)

    # Data rows, grouped by Analysis Group
    row_idx = 1
    for group, members_it in itertools.groupby(_TOC_ROWS, key=lambda r: r[2]):
        members = list(members_it)

        # Group-header spanning row (merge all 3 cells)
        grp_row = tbl.add_row()
        grp_row.cells[0].merge(grp_row.cells[1]).merge(grp_row.cells[2])
        grp_cell = grp_row.cells[0]
        grp_cell.text = group
        _set_cell_bg(grp_cell, GRP_COLOR)
        run = grp_cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)

        for i, (sheet, description, _) in enumerate(members):
            data_row = tbl.add_row()
            bg = ALT_COLOR if i % 2 == 0 else WHITE
            for cell, txt, w in zip(
                data_row.cells,
                [str(row_idx), sheet, description],
                col_widths,
                strict=False,
            ):
                cell.text = txt
                cell.width = w
                _set_cell_bg(cell, bg)
                cell.paragraphs[0].runs[0].font.size = Pt(9)
            row_idx += 1

    doc.save(path)


@task(
    kwargs={
        "depends_on": {
            # ── Reference ──────────────────────────────────────────────────────
            "lipid_list": _TAB / "lipid_data_table.pkl",
            "n_lipid": _DESC / "n_per_analysis_lipid.pkl",
            "n_prs": _DESC / "n_per_analysis_prs.pkl",
            # ── PRS ~ Subtype Probability ───────────────────────────────────────
            "prs_std": _PRS / "univariate_prs_results_standard_cov.pkl",
            # ── Lipid ~ Subtype Probability (PANSS sensitivity) ────────────────
            "lip_std": _LIP / "univariate_lipid_results.pkl",
            "lip_cov_diag": _LIP / "univariate_lipid_results_cov_diagnosis.pkl",
            "lip_cov_med": _LIP / "univariate_lipid_results_cov_med.pkl",
            "lip_cov_med_diag": _LIP / "univariate_lipid_results_cov_med_and_diag.pkl",
            "lip_cov_panss_pos": _LIP / "univariate_lipid_results_cov_panss.pkl",
            "lip_cov_panss_neg": _LIP / "univariate_lipid_results_cov_panss_neg.pkl",
            "lip_cov_panss_gen": _LIP / "univariate_lipid_results_cov_panss_gen.pkl",
            "lip_cov_panss_tot": _LIP
            / "univariate_lipid_results_cov_panss_total_score.pkl",
            # ── Lipid Class Enrichment (PANSS sensitivity) ─────────────────────
            "enrich_std": _LIP / "lipid_enrichment_results.pkl",
            "enrich_cov_diag": _LIP / "lipid_enrichment_results_cov_diagnosis.pkl",
            "enrich_cov_med": _LIP / "lipid_enrichment_results_cov_med.pkl",
            "enrich_cov_med_diag": _LIP
            / "lipid_enrichment_results_cov_med_and_diag.pkl",
            "enrich_cov_panss_pos": _LIP / "lipid_enrichment_results_cov_panss.pkl",
            "enrich_cov_panss_neg": _LIP / "lipid_enrichment_results_cov_panss_neg.pkl",
            "enrich_cov_panss_gen": _LIP / "lipid_enrichment_results_cov_panss_gen.pkl",
            "enrich_cov_panss_tot": _LIP
            / "lipid_enrichment_results_cov_panss_total_score.pkl",
            # ── Lipid Class Enrichment (medication class sensitivity) ───────────
            "enrich_cov_antidepressants": _MEDADJ
            / "lipid_enrichment_results_cov_antidepressants.pkl",
            "enrich_cov_antipsychotics": _MEDADJ
            / "lipid_enrichment_results_cov_antipsychotics.pkl",
            "enrich_cov_mood_stabilizers": _MEDADJ
            / "lipid_enrichment_results_cov_mood_stabilizers.pkl",
            "enrich_cov_tranquilizers": _MEDADJ
            / "lipid_enrichment_results_cov_tranquilizers.pkl",
            # ── PRS ~ PANSS subscales ──────────────────────────────────────────
            "prs_panss_std__pos": _PPAN
            / "univariate_prs_results_panss_standard_cov_panss_sum_pos.pkl",
            "prs_panss_std__neg": _PPAN
            / "univariate_prs_results_panss_standard_cov_panss_sum_neg.pkl",
            "prs_panss_std__gen": _PPAN
            / "univariate_prs_results_panss_standard_cov_panss_sum_gen.pkl",
            "prs_panss_std__tot": _PPAN
            / "univariate_prs_results_panss_standard_cov_panss_total_score.pkl",
            # ── Lipid (Top 20) ~ PANSS subscales ──────────────────────────────
            "lip_panss_std__pos": _LPAN
            / "univariate_lipid_results_top20_standard_panss_sum_pos.pkl",
            "lip_panss_std__neg": _LPAN
            / "univariate_lipid_results_top20_standard_panss_sum_neg.pkl",
            "lip_panss_std__gen": _LPAN
            / "univariate_lipid_results_top20_standard_panss_sum_gen.pkl",
            "lip_panss_std__tot": _LPAN
            / "univariate_lipid_results_top20_standard_panss_total_score.pkl",
            "lip_panss_med__pos": _LPAN
            / "univariate_lipid_results_top20_cov_med_panss_sum_pos.pkl",
            "lip_panss_med__neg": _LPAN
            / "univariate_lipid_results_top20_cov_med_panss_sum_neg.pkl",
            "lip_panss_med__gen": _LPAN
            / "univariate_lipid_results_top20_cov_med_panss_sum_gen.pkl",
            "lip_panss_med__tot": _LPAN
            / "univariate_lipid_results_top20_cov_med_panss_total_score.pkl",
            "lip_panss_diag__pos": _LPAN
            / "univariate_lipid_results_top20_cov_diagnosis_panss_sum_pos.pkl",
            "lip_panss_diag__neg": _LPAN
            / "univariate_lipid_results_top20_cov_diagnosis_panss_sum_neg.pkl",
            "lip_panss_diag__gen": _LPAN
            / "univariate_lipid_results_top20_cov_diagnosis_panss_sum_gen.pkl",
            "lip_panss_diag__tot": _LPAN
            / "univariate_lipid_results_top20_cov_diagnosis_panss_total_score.pkl",
            "lip_panss_med_diag__pos": _LPAN
            / "univariate_lipid_results_top20_cov_med_and_diag_panss_sum_pos.pkl",
            "lip_panss_med_diag__neg": _LPAN
            / "univariate_lipid_results_top20_cov_med_and_diag_panss_sum_neg.pkl",
            "lip_panss_med_diag__gen": _LPAN
            / "univariate_lipid_results_top20_cov_med_and_diag_panss_sum_gen.pkl",
            "lip_panss_med_diag__tot": _LPAN
            / "univariate_lipid_results_top20_cov_med_and_diag_panss_total_score.pkl",
            # ── Lipid Enrichment ~ PANSS subscales ────────────────────────────
            "enrich_panss_std__pos": _LPAN
            / "lipid_enrichment_results_panss_sum_pos.pkl",
            "enrich_panss_std__neg": _LPAN
            / "lipid_enrichment_results_panss_sum_neg.pkl",
            "enrich_panss_std__gen": _LPAN / "lipid_enrichment_results_panss_gen.pkl",
            "enrich_panss_std__tot": _LPAN
            / "lipid_enrichment_results_panss_total_score.pkl",
            # ── Incremental R² ─────────────────────────────────────────────────
            "r2_incremental": _R2 / "incremental_r2_table.pkl",
            "r2_individual": _R2 / "individual_predictor_r2.pkl",
            # ── Mediation ──────────────────────────────────────────────────────
            "mediation": _MED / "mediation_analysis_results.pkl",
            # ── CCA ────────────────────────────────────────────────────────────
            "cca_results": BLD_RESULTS / "cca_regression" / "results.pkl",
        },
        "produces": {
            "excel": BLD_RESULTS / "supplementary_tables" / "supplementary_tables.xlsx",
            "toc_docx": BLD_RESULTS / "supplementary_tables" / "supplementary_toc.docx",
        },
    }
)
def task_export_all_supplementary_tables(depends_on, produces):
    d = depends_on  # shorthand

    def load(key: str) -> pd.DataFrame:
        return pd.read_pickle(d[key])

    def stack_panss(label_key_pairs: list[tuple[str, str]]) -> pd.DataFrame:
        """Load each file, prepend a 'panss_outcome' column, and concatenate."""
        frames = []
        for label, key in label_key_pairs:
            df = load(key).copy()
            df.insert(0, "panss_outcome", label)
            frames.append(df)
        return pd.concat(frames)

    with pd.ExcelWriter(produces["excel"], engine="openpyxl") as writer:
        # ── Sheet 01: Table of Contents ────────────────────────────────────────
        toc_df = pd.DataFrame(
            _TOC_ROWS,
            columns=["Sheet", "Description", "Analysis Group"],
        )
        toc_df.to_excel(writer, sheet_name="01_TOC", index=False)

        # ── Reference ──────────────────────────────────────────────────────────
        load("lipid_list").to_excel(writer, sheet_name="02_lipid_list")
        load("n_lipid").to_excel(writer, sheet_name="03_n_lipid")
        load("n_prs").to_excel(writer, sheet_name="04_n_prs")

        # ── PRS ~ Subtype Probability ───────────────────────────────────────────
        load("prs_std").to_excel(writer, sheet_name="05_prs_std")

        # ── Lipid ~ Subtype Probability (PANSS sensitivity) ────────────────────
        load("lip_std").to_excel(writer, sheet_name="06_lip_std")
        load("lip_cov_diag").to_excel(writer, sheet_name="07_lip_cov_diag")
        load("lip_cov_med").to_excel(writer, sheet_name="08_lip_cov_med")
        load("lip_cov_med_diag").to_excel(writer, sheet_name="09_lip_cov_med_diag")
        load("lip_cov_panss_pos").to_excel(writer, sheet_name="10_lip_cov_panss_pos")
        load("lip_cov_panss_neg").to_excel(writer, sheet_name="11_lip_cov_panss_neg")
        load("lip_cov_panss_gen").to_excel(writer, sheet_name="12_lip_cov_panss_gen")
        load("lip_cov_panss_tot").to_excel(writer, sheet_name="13_lip_cov_panss_tot")

        # ── Lipid Class Enrichment (PANSS sensitivity) ─────────────────────────
        load("enrich_std").to_excel(writer, sheet_name="14_enrich_std")
        load("enrich_cov_diag").to_excel(writer, sheet_name="15_enrich_cov_diag")
        load("enrich_cov_med").to_excel(writer, sheet_name="16_enrich_cov_med")
        load("enrich_cov_med_diag").to_excel(
            writer, sheet_name="17_enrich_cov_med_diag"
        )
        load("enrich_cov_panss_pos").to_excel(
            writer, sheet_name="18_enrich_cov_panss_pos"
        )
        load("enrich_cov_panss_neg").to_excel(
            writer, sheet_name="19_enrich_cov_panss_neg"
        )
        load("enrich_cov_panss_gen").to_excel(
            writer, sheet_name="20_enrich_cov_panss_gen"
        )
        load("enrich_cov_panss_tot").to_excel(
            writer, sheet_name="21_enrich_cov_panss_tot"
        )

        # ── Lipid Class Enrichment (medication class sensitivity) ──────────────
        load("enrich_cov_antidepressants").to_excel(
            writer, sheet_name="22_enrich_cov_antidepressants"
        )
        load("enrich_cov_antipsychotics").to_excel(
            writer, sheet_name="23_enrich_cov_antipsychotics"
        )
        load("enrich_cov_mood_stabilizers").to_excel(
            writer, sheet_name="24_enrich_cov_mood_stabilizers"
        )
        load("enrich_cov_tranquilizers").to_excel(
            writer, sheet_name="25_enrich_cov_tranquilizers"
        )

        # ── PRS ~ PANSS subscales (stacked) ────────────────────────────────────
        _PANSS_LABELS = [
            ("PANSS Positive", "__pos"),
            ("PANSS Negative", "__neg"),
            ("PANSS General", "__gen"),
            ("PANSS Total Score", "__tot"),
        ]
        stack_panss(
            [(lbl, f"prs_panss_std{sfx}") for lbl, sfx in _PANSS_LABELS]
        ).to_excel(writer, sheet_name="26_prs_panss_std")

        # ── Lipid (Top 20) ~ PANSS subscales (stacked) ─────────────────────────
        stack_panss(
            [(lbl, f"lip_panss_std{sfx}") for lbl, sfx in _PANSS_LABELS]
        ).to_excel(writer, sheet_name="27_lip_panss_std")
        stack_panss(
            [(lbl, f"lip_panss_med{sfx}") for lbl, sfx in _PANSS_LABELS]
        ).to_excel(writer, sheet_name="28_lip_panss_cov_med")
        stack_panss(
            [(lbl, f"lip_panss_diag{sfx}") for lbl, sfx in _PANSS_LABELS]
        ).to_excel(writer, sheet_name="29_lip_panss_cov_diag")
        stack_panss(
            [(lbl, f"lip_panss_med_diag{sfx}") for lbl, sfx in _PANSS_LABELS]
        ).to_excel(writer, sheet_name="30_lip_panss_cov_med_diag")

        # ── Lipid Enrichment ~ PANSS subscales (stacked) ───────────────────────
        stack_panss(
            [
                ("PANSS Positive", "enrich_panss_std__pos"),
                ("PANSS Negative", "enrich_panss_std__neg"),
                ("PANSS General", "enrich_panss_std__gen"),
                ("PANSS Total Score", "enrich_panss_std__tot"),
            ]
        ).to_excel(writer, sheet_name="31_enrich_panss_std")

        # ── Incremental R² ─────────────────────────────────────────────────────
        load("r2_incremental").to_excel(writer, sheet_name="32_r2_incremental")
        load("r2_individual").to_excel(writer, sheet_name="33_r2_individual")

        # ── Mediation ──────────────────────────────────────────────────────────
        load("mediation").to_excel(writer, sheet_name="34_mediation")

        # ── CCA canonical loadings ──────────────────────────────────────────────
        cca_res = load("cca_results")
        prs_df = (
            cca_res["prs_loadings"]
            .rename("loading")
            .rename_axis("variable")
            .reset_index()
        )
        prs_df.insert(0, "block", "PRS")
        lip_df = (
            cca_res["lipid_class_loadings"]
            .rename("loading")
            .rename_axis("variable")
            .reset_index()
        )
        lip_df.insert(0, "block", "Lipid class")
        cca_df = pd.concat([prs_df, lip_df], ignore_index=True)
        cca_df.to_excel(writer, sheet_name="35_cca_loadings", index=False)

    # ── Write TOC docx ─────────────────────────────────────────────────────────
    _write_toc_docx(produces["toc_docx"])
